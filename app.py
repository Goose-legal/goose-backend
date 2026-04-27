from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from io import BytesIO
from datetime import datetime
import anthropic
import os

app = Flask(__name__)
CORS(app)

client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

@app.route("/health")
def health():
    return "ok"

@app.route("/analyse", methods=["POST"])
def analyse():
    data = request.json
    case_text = data.get("caseText", "")

    if not case_text:
        return jsonify({"error": "Ingen text skickades"}), 400

    def generate():
        with client.messages.stream(
            model="claude-haiku-4-5-20251001",
            max_tokens=1500,
            messages=[{
                "role": "user",
                "content": f"""Du ska analysera ett svenskt rattsfall och endast redovisa Hogsta domstolens avgörande.

Viktigt:
- Bygg endast pa det som uttryckligen framgar av texten.
- Sarskild noga mellan bakgrund, rattsfrage, HD:s motivering och prejudikatverkan.
- Ta inte med egna antaganden.
- Om information saknas eller ar osakar, skriv: Framgår inte tydligt av texten.
- Efter varje pastand, citera det relevanta stycket fran rattsfallet inom citationstecken.

Svara med foljande rubriker och inget annat:

HD:S BESLUT
Beskriv kort vad HD kom fram till och hur malet avgjordes.

RÄTTSFRAGA
Forklara vilken rattslig huvudfraga HD provade och varfor fragan var juridiskt viktig.

DOMSKÄL
Beskriv steg for steg hur HD resonerade for att na sitt avgörande.

LEGALA PRINCIPER
Ange vilka rattregler eller principer som HD slog fast.

SKILJAKTIG MENING
Ange om nagon ledamot var skiljaktig och vad de ansag.

PREJUDIKAT
Forklara vilket vagledande varde avgörandet kan fa.

Krav:
- Max 2 mening per rubrik
- Max 400 ord totalt
- Enkel och tydlig svenska
- Inga punktlistor
- Ingen markdown

Text:
<rattsfall>
{case_text}
</rattsfall>"""
            }]
        ) as stream:
            for text in stream.text_stream:
                yield text

    return app.response_class(generate(), mimetype='text/plain')


@app.route("/download", methods=["POST"])
def download():
    data = request.json
    analysis = data.get("analysis", "")
    case_name = data.get("caseName", "goose-analys")

    doc = DocxDocument()

    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.2)
    section.left_margin = Inches(1.4)
    section.right_margin = Inches(1.4)

    # Titel
    title = doc.add_paragraph()
    title_run = title.add_run("Goose")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.alignment = 0

    # Datum
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Analys genererad {datetime.now().strftime('%d %B %Y')}")
    date_run.font.name = "Arial"
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    date_para.alignment = 0

    headings = ["HD:S BESLUT", "RÄTTSFRAGA", "RÄTTSFRÅGA", "DOMSKÄL", "LEGALA PRINCIPER", "SKILJAKTIG MENING", "PREJUDIKAT"]

    # Dela upp analysen vid rubriker
    import re
    pattern = '(' + '|'.join(re.escape(h) for h in headings) + ')'
    parts = re.split(pattern, analysis)

    for part in parts:
        part = part.strip()
        if not part:
            continue

        if part in headings:
            # Tomt stycke innan rubrik
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
        else:
            # Brödtext — dela vid meningar för bättre läsbarhet
            sentences = part.replace('. "', '.\n"').replace('." ', '."\n').split('\n')
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(sentence)
                run.font.name = "Georgia"
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = Pt(15)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = "".join(c for c in case_name if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"{safe_name}.docx" if safe_name else "goose-analys.docx"

    return app.response_class(
        buffer.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )

@app.route("/privacy")
def privacy():
    return """
    <html>
    <head><title>Goose - Integritetspolicy</title></head>
    <body style="font-family: Arial; max-width: 800px; margin: 40px auto; padding: 20px;">
        <h1>Integritetspolicy for Goose</h1>
        <p>Senast uppdaterad: 6 april 2026</p>
        <h2>Vilken information samlar vi in?</h2>
        <p>Goose samlar inte in personuppgifter. Den text du analyserar skickas till Anthropics API for behandling och lagras inte av oss.</p>
        <h2>Hur anvander vi informationen?</h2>
        <p>Texten du skickar in anvands enbart for att generera en analys via Anthropic Claude AI.</p>
        <h2>Kontakt</h2>
        <p>Fragor? Kontakta oss pa: alastaircs@gmail.com</p>
    </body>
    </html>
    """

if __name__ == "__main__":
    app.run(debug=True)